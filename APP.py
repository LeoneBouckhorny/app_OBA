import streamlit as st
from docx import Document
from collections import defaultdict
import os

# === Funções auxiliares ===
def formatar_texto(texto, maiusculo_estado=False):
    texto = ' '.join(texto.strip().split())
    if maiusculo_estado:
        return texto.upper()
    return ' '.join(w.capitalize() for w in texto.split())

def processar_docx(uploaded_file):
    doc = Document(uploaded_file)
    dados = []

    for tabela in doc.tables:
        for i, linha in enumerate(tabela.rows):
            if i == 0:
                continue  # Ignora cabeçalho
            valores = [c.text.strip() for c in linha.cells]
            if len(valores) >= 8:  # Garantir que todas as colunas existam
                medalha, valido, equipe, funcao, escola, cidade, estado, nome = valores[:8]
                dados.append({
                    "Medalha": medalha,
                    "Valido": valido,
                    "Equipe": equipe,
                    "Funcao": funcao.lower(),
                    "Escola": escola,
                    "Cidade": cidade,
                    "Estado": estado,
                    "Nome": nome
                })

    equipes = defaultdict(list)
    for item in dados:
        equipes[item["Equipe"]].append(item)

    # Ordenar equipes pelo lançamento válido do primeiro membro (em ordem crescente)
    def valor_valido(membros):
        try:
            return float(membros[0]['Valido'].replace(',', '.'))
        except:
            return float('inf')  # Valores vazios vão para o final

    equipes_ordenadas = sorted(equipes.items(), key=lambda x: valor_valido(x[1]))

    novo_doc = Document()
    for equipe, membros in equipes_ordenadas:
        lider = [m for m in membros if "líder" in m["Funcao"] or "lider" in m["Funcao"]]
        acompanhante = [m for m in membros if "acompanhante" in m["Funcao"]]
        alunos = [m for m in membros if "aluno" in m["Funcao"]]
        alunos_sorted = sorted(alunos, key=lambda m: formatar_texto(m["Nome"]))

        ordem_final = lider + acompanhante + alunos_sorted

        for membro in ordem_final:
            novo_doc.add_paragraph(formatar_texto(membro["Nome"]))

        if membros:
            novo_doc.add_paragraph(f"Equipe: {equipe.split()[-1]}")
            novo_doc.add_paragraph(f"Lançamentos Válidos: {membros[0]['Valido']} m")
            novo_doc.add_paragraph(formatar_texto(membros[0]["Escola"]))
            novo_doc.add_paragraph(f"{formatar_texto(membros[0]['Cidade'])} / {formatar_texto(membros[0]['Estado'], maiusculo_estado=True)}")
        novo_doc.add_paragraph("")  # Separação entre equipes

    return novo_doc

# === Interface Streamlit ===
st.title("🏅 Organizador de Equipes e Resultados")
st.write("Faça upload do arquivo `.docx` e baixe o arquivo formatado com lançamentos válidos, ordenado por valor crescente.")

uploaded_file = st.file_uploader("Envie o arquivo DOCX", type=["docx"])

if uploaded_file:
    # Nome base do arquivo original (sem extensão)
    nome_base = os.path.splitext(uploaded_file.name)[0]
    novo_nome = f"{nome_base}_FORMATADO.docx"

    # Processar o arquivo
    novo_doc = processar_docx(uploaded_file)

    # Prévia (usar st.code para evitar erros de renderização)
    st.subheader("Prévia das primeiras equipes:")
    preview = [p.text for p in novo_doc.paragraphs[:20]]
    st.code("\n".join(preview), language="text")

    # Salvar e disponibilizar para download
    novo_doc.save(novo_nome)
    with open(novo_nome, "rb") as f:
        st.download_button(
            label=f"📥 Baixar {novo_nome}",
            data=f,
            file_name=novo_nome,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
